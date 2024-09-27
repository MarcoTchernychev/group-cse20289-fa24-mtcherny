#Marco Tchernychev
#mtcherny@nd.edu
import argparse
import docx

#INPUT: a docx object and a dictionary
#OUTPUT: the docx object wth a table inserted that comes from dict
#PURPOSE: add a table to the doc using the inputted dict
def insertTable(doc, dict):
    table = doc.add_table(rows=len(dict), cols=2)
    table.style = 'Table Grid'
    for i, key in enumerate(dict.keys()):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = str(dict[key])
    return doc
#INPUT: a text file, dictionary, png file, and the name of an output docx file
#OUTPUT: none, just saves the docx file
#PURPOSE: add the text to the docx file, add the table to the docx file, add the image to the docs file
def makeReport(textfilepath, dict, pngfilepath, outputrprt):
    doc = docx.Document()
    with open(textfilepath, 'r') as file:
        contents = file.read()
    doc.add_paragraph(contents)
    insertTable(doc, dict)
    doc.add_picture(pngfilepath)
    doc.save(outputrprt)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("txtfilepath", type=str)
    parser.add_argument("pngfilepath", type=str)
    parser.add_argument("outputrprt", type=str)
    args = parser.parse_args()
    txtfilepath = args.txtfilepath
    pngfilepath = args.pngfilepath
    outputrprt = args.outputrprt
    dummydict = {"hey":0, "ho":0, "let's":0, "go":0}
    makeReport(txtfilepath, dummydict, pngfilepath, outputrprt)

